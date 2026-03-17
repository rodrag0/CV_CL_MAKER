from __future__ import annotations

import json
import os
import shutil
import subprocess
import zipfile
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.tailor import TailoredApplication


APP_ROOT = Path(__file__).resolve().parents[1]
GENERATED_ROOT = Path(os.environ.get("APP_GENERATED_DIR", APP_ROOT / "generated"))

ACCENT = RGBColor(46, 78, 62)
DARK = RGBColor(34, 34, 34)
MUTED = RGBColor(90, 90, 90)


@dataclass(frozen=True)
class GeneratedFiles:
    pack_id: str
    job_id: str
    output_root: Path
    ats_docx: Path
    designed_docx: Path
    letter_docx: Path | None
    ats_pdf: Path | None
    designed_pdf: Path | None
    letter_pdf: Path | None
    bundle_zip: Path
    manifest_path: Path
    job_posting_path: Path


def cleanup_old_packs(max_age_hours: int = 24) -> None:
    GENERATED_ROOT.mkdir(parents=True, exist_ok=True)
    cutoff = datetime.now(timezone.utc) - timedelta(hours=max_age_hours)
    for child in GENERATED_ROOT.iterdir():
        try:
            modified = datetime.fromtimestamp(child.stat().st_mtime, tz=timezone.utc)
        except OSError:
            continue
        if modified < cutoff:
            if child.is_dir():
                shutil.rmtree(child, ignore_errors=True)
            else:
                child.unlink(missing_ok=True)


def build_output_paths(pack_root: Path, job_id: str) -> tuple[Path, Path, Path, Path]:
    cv_job_dir = pack_root / "CV" / job_id
    cover_job_dir = pack_root / "Cover Letters" / job_id
    cv_files_dir = cv_job_dir / "files"
    cover_files_dir = cover_job_dir / "files"
    cv_files_dir.mkdir(parents=True, exist_ok=True)
    cover_files_dir.mkdir(parents=True, exist_ok=True)
    return cv_job_dir, cover_job_dir, cv_files_dir, cover_files_dir


def set_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)


def ptext(paragraph, text: str, *, bold: bool = False, size: float = 10.5, color=DARK, italic: bool = False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


def add_divider(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2F4F4F")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_cell_shading(cell, color_hex: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_table_no_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil")
        borders.append(elem)
    tbl_pr.append(borders)


def heading(doc: Document, label: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)
    ptext(paragraph, label.upper(), bold=True, size=11, color=ACCENT)
    add_divider(paragraph)


def designed_heading(cell, label: str) -> None:
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    ptext(paragraph, label.upper(), bold=True, size=10.5, color=ACCENT)


def write_ats_cv(application: TailoredApplication, output_docx: Path) -> None:
    doc = Document()
    set_page(doc)

    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ptext(header, application.profile.name, bold=True, size=23, color=ACCENT)

    role = doc.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ptext(role, application.role_label, bold=True, size=11, color=DARK)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ptext(contact, application.contact_line, size=10, color=MUTED)

    links = doc.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ptext(links, application.links_line, size=10, color=ACCENT)

    heading(doc, "Professional Summary" if application.language == "en" else "Profil")
    summary = doc.add_paragraph()
    summary.paragraph_format.space_after = Pt(0)
    ptext(summary, application.summary, size=10.5)

    heading(doc, "Core Skills" if application.language == "en" else "Kernkompetenzen")
    for item in application.skills:
        bullet = doc.add_paragraph(style="List Bullet")
        bullet.paragraph_format.space_after = Pt(1)
        ptext(bullet, item, size=10.5)

    heading(doc, "Professional Experience" if application.language == "en" else "Berufserfahrung")
    for experience in application.experiences:
        company = doc.add_paragraph()
        ptext(company, experience.company, bold=True, size=11)
        title = doc.add_paragraph()
        ptext(title, experience.title, bold=True, size=10.5, color=ACCENT)
        ptext(title, " | ", size=10.5, color=MUTED)
        ptext(title, experience.date, italic=True, size=10.5, color=MUTED)
        for item in experience.bullets:
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.paragraph_format.space_after = Pt(1)
            ptext(bullet, item, size=10.5)

    heading(doc, "Education" if application.language == "en" else "Ausbildung")
    for item in application.education:
        bullet = doc.add_paragraph(style="List Bullet")
        bullet.paragraph_format.space_after = Pt(1)
        ptext(bullet, item, size=10.5)

    heading(doc, "Selected Projects" if application.language == "en" else "Ausgewaehlte Projekte")
    for item in application.projects:
        bullet = doc.add_paragraph(style="List Bullet")
        bullet.paragraph_format.space_after = Pt(1)
        ptext(bullet, item, size=10.5)

    if application.founder_experience:
        heading(
            doc,
            "Entrepreneurial Experience" if application.language == "en" else "Unternehmerische Erfahrung",
        )
        for item in application.founder_experience:
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.paragraph_format.space_after = Pt(1)
            ptext(bullet, item, size=10.5)

    doc.save(output_docx)


def write_designed_cv(application: TailoredApplication, output_docx: Path) -> None:
    doc = Document()
    set_page(doc)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_no_borders(table)

    left, right = table.rows[0].cells
    left.width = Inches(2.15)
    right.width = Inches(4.95)
    set_cell_shading(left, "F3F6F8")

    profile = application.profile

    name = left.add_paragraph()
    name.paragraph_format.space_after = Pt(2)
    ptext(name, profile.name, bold=True, size=17, color=ACCENT)

    role = left.add_paragraph()
    ptext(role, profile.headline.replace(" | ", "\n"), bold=True, size=10, color=DARK)

    designed_heading(left, "Contact")
    for line in [profile.city, profile.address, profile.phone, profile.email]:
        paragraph = left.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        ptext(paragraph, line, size=9.5, color=DARK)

    designed_heading(left, "Links")
    for line in [profile.github, profile.linkedin]:
        paragraph = left.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        ptext(paragraph, line, size=9.5, color=ACCENT)

    designed_heading(left, "Core Skills" if application.language == "en" else "Kernkompetenzen")
    for item in application.skills:
        bullet = left.add_paragraph(style="List Bullet")
        bullet.paragraph_format.space_after = Pt(0)
        ptext(bullet, item, size=9.1, color=DARK)

    designed_heading(left, "Languages" if application.language == "en" else "Sprachen")
    for line in profile.languages:
        bullet = left.add_paragraph(style="List Bullet")
        bullet.paragraph_format.space_after = Pt(0)
        ptext(bullet, line, size=9.1, color=DARK)

    title = right.add_paragraph()
    ptext(title, application.role_label, bold=True, size=14, color=DARK)

    profile_head = right.add_paragraph()
    ptext(profile_head, "PROFILE" if application.language == "en" else "PROFIL", bold=True, size=10.5, color=ACCENT)
    summary = right.add_paragraph()
    summary.paragraph_format.space_after = Pt(6)
    ptext(summary, application.summary, size=10.2)

    experience_head = right.add_paragraph()
    ptext(experience_head, "EXPERIENCE" if application.language == "en" else "ERFAHRUNG", bold=True, size=10.5, color=ACCENT)
    for experience in application.experiences:
        title_line = right.add_paragraph()
        ptext(title_line, f"{experience.title} | {experience.company}", bold=True, size=10.2)
        date_line = right.add_paragraph()
        date_line.paragraph_format.space_after = Pt(1)
        ptext(date_line, experience.date, italic=True, size=9.5, color=MUTED)
        for item in experience.bullets:
            bullet = right.add_paragraph(style="List Bullet")
            bullet.paragraph_format.space_after = Pt(0)
            ptext(bullet, item, size=9.8)

    education_head = right.add_paragraph()
    education_head.paragraph_format.space_before = Pt(7)
    ptext(education_head, "EDUCATION" if application.language == "en" else "AUSBILDUNG", bold=True, size=10.5, color=ACCENT)
    for item in application.education:
        bullet = right.add_paragraph(style="List Bullet")
        bullet.paragraph_format.space_after = Pt(0)
        ptext(bullet, item, size=9.8)

    projects_head = right.add_paragraph()
    projects_head.paragraph_format.space_before = Pt(7)
    ptext(projects_head, "SELECTED PROJECTS" if application.language == "en" else "AUSGEWAEHLTE PROJEKTE", bold=True, size=10.5, color=ACCENT)
    for item in application.projects:
        bullet = right.add_paragraph(style="List Bullet")
        bullet.paragraph_format.space_after = Pt(0)
        ptext(bullet, item, size=9.8)

    if application.founder_experience:
        founder_head = right.add_paragraph()
        founder_head.paragraph_format.space_before = Pt(7)
        label = "ENTREPRENEURIAL EXPERIENCE" if application.language == "en" else "UNTERNEHMERISCHE ERFAHRUNG"
        ptext(founder_head, label, bold=True, size=10.5, color=ACCENT)
        for item in application.founder_experience:
            bullet = right.add_paragraph(style="List Bullet")
            bullet.paragraph_format.space_after = Pt(0)
            ptext(bullet, item, size=9.8)

    doc.save(output_docx)


def write_cover_letter(application: TailoredApplication, output_docx: Path) -> None:
    doc = Document()
    set_page(doc)

    profile = application.profile

    name = doc.add_paragraph()
    ptext(name, profile.name, bold=True, size=20, color=ACCENT)

    sub = doc.add_paragraph()
    ptext(sub, profile.headline, size=10.5, color=MUTED)

    contact = doc.add_paragraph()
    ptext(contact, application.contact_line, size=10, color=DARK)

    links = doc.add_paragraph()
    ptext(links, application.links_line, size=10, color=ACCENT)
    add_divider(links)

    doc.add_paragraph()
    recipient = doc.add_paragraph()
    ptext(recipient, application.cover_letter_recipient, size=11)

    date_line = doc.add_paragraph()
    date_line.paragraph_format.space_before = Pt(8)
    ptext(date_line, application.cover_letter_date, size=11)

    subject = doc.add_paragraph()
    subject.paragraph_format.space_before = Pt(8)
    ptext(subject, application.cover_letter_subject, bold=True, size=11.3, color=ACCENT)

    for index, paragraph in enumerate(application.cover_letter_body):
        body = doc.add_paragraph()
        body.paragraph_format.space_before = Pt(7 if index else 10)
        body.paragraph_format.space_after = Pt(0)
        ptext(body, paragraph, size=11)

    doc.save(output_docx)


def try_convert_to_pdf(docx_path: Path) -> Path | None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        common_paths = (
            Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
            Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
        )
        for candidate in common_paths:
            if candidate.exists():
                soffice = str(candidate)
                break
    if not soffice:
        return None
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
            check=True,
            capture_output=True,
            text=True,
        )
    except (OSError, subprocess.CalledProcessError):
        return None
    pdf_path = docx_path.with_suffix(".pdf")
    return pdf_path if pdf_path.exists() else None


def write_manifest(application: TailoredApplication, generated: GeneratedFiles) -> None:
    payload = {
        "pack_id": generated.pack_id,
        "job_id": application.job_id,
        "company": application.company,
        "role": application.role,
        "language": application.language,
        "tailoring_mode": application.tailoring_mode,
        "ai_model": application.ai_model or None,
        "focuses": list(application.focuses),
        "generated_files": {
            "ats_docx": str(generated.ats_docx.relative_to(generated.output_root)),
            "designed_docx": str(generated.designed_docx.relative_to(generated.output_root)),
            "letter_docx": str(generated.letter_docx.relative_to(generated.output_root)) if generated.letter_docx else None,
            "ats_pdf": str(generated.ats_pdf.relative_to(generated.output_root)) if generated.ats_pdf else None,
            "designed_pdf": str(generated.designed_pdf.relative_to(generated.output_root)) if generated.designed_pdf else None,
            "letter_pdf": str(generated.letter_pdf.relative_to(generated.output_root)) if generated.letter_pdf else None,
            "bundle_zip": str(generated.bundle_zip.relative_to(generated.output_root)),
        },
    }
    generated.manifest_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def build_bundle_zip(generated: GeneratedFiles) -> None:
    files_to_add = [
        generated.ats_docx,
        generated.designed_docx,
        generated.manifest_path,
        generated.job_posting_path,
    ]
    optional_files = [
        generated.letter_docx,
        generated.ats_pdf,
        generated.designed_pdf,
        generated.letter_pdf,
    ]
    files_to_add.extend(path for path in optional_files if path is not None)

    with zipfile.ZipFile(generated.bundle_zip, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in files_to_add:
            archive.write(path, arcname=str(path.relative_to(generated.output_root)))


def generate_application_pack(
    application: TailoredApplication,
    *,
    pack_id: str,
    create_cover_letter: bool = True,
) -> GeneratedFiles:
    cleanup_old_packs()
    pack_root = GENERATED_ROOT / pack_id
    pack_root.mkdir(parents=True, exist_ok=True)

    cv_job_dir, cover_job_dir, cv_files_dir, cover_files_dir = build_output_paths(pack_root, application.job_id)

    ats_docx = cv_files_dir / "Rodrigo_Ponce_Cortes_CV_ATS_Styled.docx"
    designed_docx = cv_files_dir / "Rodrigo_Ponce_Cortes_CV_Designed.docx"
    letter_docx = cover_files_dir / "Rodrigo_Ponce_Cortes_Cover_Letter_Styled.docx" if create_cover_letter else None

    write_ats_cv(application, ats_docx)
    write_designed_cv(application, designed_docx)
    if letter_docx:
        write_cover_letter(application, letter_docx)

    ats_pdf = try_convert_to_pdf(ats_docx)
    designed_pdf = try_convert_to_pdf(designed_docx)
    letter_pdf = try_convert_to_pdf(letter_docx) if letter_docx else None

    job_posting_path = pack_root / "job_posting.txt"
    job_posting_path.write_text(application.job_posting.strip() + "\n", encoding="utf-8")
    manifest_path = pack_root / "application_manifest.json"
    bundle_zip = pack_root / f"{application.job_id}_application_pack.zip"

    generated = GeneratedFiles(
        pack_id=pack_id,
        job_id=application.job_id,
        output_root=pack_root,
        ats_docx=ats_docx,
        designed_docx=designed_docx,
        letter_docx=letter_docx,
        ats_pdf=ats_pdf,
        designed_pdf=designed_pdf,
        letter_pdf=letter_pdf,
        bundle_zip=bundle_zip,
        manifest_path=manifest_path,
        job_posting_path=job_posting_path,
    )
    write_manifest(application, generated)
    build_bundle_zip(generated)
    return generated
